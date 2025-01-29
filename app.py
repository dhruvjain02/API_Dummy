from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import io
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'  # Change this to a secure secret key

# Configuration
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    """Render the landing page."""
    try:
        return render_template('index.html')
    except Exception as e:
        logger.error(f"Error rendering index page: {str(e)}")
        return "An error occurred while loading the page", 500

@app.route('/static/css/<path:filename>')
def serve_css(filename):
    return send_from_directory('static/css', filename)

@app.route('/form')
def form():
    """Render the form page."""
    try:
        return render_template('form.html')
    except Exception as e:
        logger.error(f"Error rendering form page: {str(e)}")
        return "An error occurred while loading the form", 500

@app.route('/thankyou')
def thankyou():
    """Render the thank you page."""
    try:
        return render_template('thankyou.html')
    except Exception as e:
        logger.error(f"Error rendering thank you page: {str(e)}")
        return "An error occurred while loading the thank you page", 500

@app.route('/submit', methods=['POST'])
def submit():
    """Handle form submission and generate report."""
    try:
        # Extract form data with validation
        form_data = {
            "Insert Case Number": request.form.get('case_number', "N/A"),
            "Insert Date": request.form.get('date_of_report', datetime.now().strftime("%Y-%m-%d")),
            "Insert Name and Title": request.form.get('report_prepared_by', "N/A"),
            "Insert Device Model": request.form.get('model', "N/A"),
            "Insert Color": request.form.get('color', "N/A"),
            "Safety Glass Yes or No": request.form.get('safety_glass', "N/A"),
            "Back Cover Yes or No": request.form.get('back_cover', "N/A"),
            "Insert RAM": request.form.get('ram', "N/A"),
            "Insert Internal Memory": request.form.get('internal_memory', "N/A"),
            "Insert Camera Details": request.form.get('camera_specs', "N/A"),
            "Insert Cameras Check": request.form.get('cameras_check', "N/A"),
            "Insert Battery Percentage": request.form.get('battery_percentage', "N/A"),
            "Insert SIM Slots": request.form.get('sim_slots', "N/A"),
            "Insert SIM Provider": request.form.get('sim_provider', "N/A"),
            "Insert Wi-Fi Status": request.form.get('wifi_connected', "N/A"),
            "Insert Bluetooth Status": request.form.get('bluetooth_status', "N/A"),
            "Insert SD Card Present": request.form.get('sd_card', "N/A"),
            "Insert SD Capacity": request.form.get('sd_capacity', "N/A"),
            "Insert Mobile Uptime": request.form.get('mobile_uptime', "N/A"),
            "Insert Time Zone": request.form.get('time_zone', "N/A"),
            "Insert Language": request.form.get('language', "N/A"),
            "Insert Installed Apps": request.form.get('installed_apps', "N/A"),
            "Insert Geo Location": request.form.get('geo_location', "N/A"),
            "Insert Build Number": request.form.get('build_no', "N/A"),
            "Insert Kernel Version": request.form.get('kernel_version', "N/A"),
            "Insert ICCID": request.form.get('iccid', "N/A"),
            "Insert IMSI": request.form.get('imsi', "N/A"),
            "Insert MEID": request.form.get('meid', "N/A"),
            "Insert Airplane Mode": request.form.get('airplane_mode', "N/A")
        }

        # Validate required fields
        required_fields = ['case_number', 'report_prepared_by', 'model']
        for field in required_fields:
            if not request.form.get(field):
                return jsonify({'error': f'Missing required field: {field}'}), 400

        # Load the Word template
        template_path = os.path.join(os.getcwd(), 'Report Format.docx')
        if not os.path.exists(template_path):
            logger.error("Report template not found")
            return jsonify({'error': 'Report template not found'}), 500

        document = Document(template_path)

        # Replace placeholders in paragraphs
        for para in document.paragraphs:
            for key, value in form_data.items():
                placeholder = f"[{key}]"
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))

        # Replace placeholders in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in form_data.items():
                        placeholder = f"[{key}]"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        # Handle device images
        angle_mapping = {
            'image_0': '[0]',
            'image_1': '[30]',
            'image_2': '[60]',
            'image_3': '[90]',
            'image_4': '[240]',
            'image_5': '[270]',
            'image_6': '[360]'
        }

        # Process and insert images
        for image_key, placeholder in angle_mapping.items():
            if image_key in request.files:
                image_file = request.files[image_key]
                if image_file and image_file.filename and allowed_file(image_file.filename):
                    try:
                        logger.info(f"Processing {image_key} with placeholder {placeholder}")
                        
                        # Find and replace image in tables
                        for table in document.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if placeholder in cell.text:
                                        # Clear existing content
                                        cell.text = ""
                                        
                                        # Add image
                                        paragraph = cell.paragraphs[0]
                                        run = paragraph.add_run()
                                        image_file.seek(0)
                                        image_data = io.BytesIO(image_file.read())
                                        run.add_picture(image_data, width=Inches(2.5))
                                        logger.info(f"Successfully inserted {image_key}")
                    except Exception as e:
                        logger.error(f"Error processing image {image_key}: {str(e)}")
                        continue

        # Save document to buffer
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f'report_{form_data["Insert Case Number"]}_{timestamp}.docx'

        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"Error generating report: {str(e)}")
        return jsonify({'error': 'An error occurred while generating the report'}), 500

@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle file size exceeded error."""
    return jsonify({'error': 'File size exceeded maximum limit'}), 413

@app.errorhandler(404)
def not_found_error(error):
    """Handle 404 errors."""
    return jsonify({'error': 'Resource not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors."""
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == '__main__':
    # Ensure templates directory exists
    if not os.path.exists('templates'):
        os.makedirs('templates')
        logger.warning("Created missing templates directory")
    
    # Ensure static directory exists
    if not os.path.exists('static/images'):
        os.makedirs('static/images')
        logger.warning("Created missing static/images directory")
    
    # Start the application
    app.run(debug=True, host='0.0.0.0', port=5000)
